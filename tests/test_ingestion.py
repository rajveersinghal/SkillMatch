import unittest
from unittest.mock import MagicMock, patch
import os
import docx

# Import modules to test
# Import modules to test
from core.ingestion import extract_text_from_docx
import core.database as database


class TestPhase4(unittest.TestCase):
    def test_docx_extraction(self):
        # Create a dummy docx
        filename = "test_resume.docx"
        doc = docx.Document()
        doc.add_paragraph("Hello World")
        doc.add_paragraph("This is a resume.")
        doc.save(filename)

        # Test extraction
        # We need a file-like object or path. ingestion.py expects variable args depending on library
        # extract_text_from_docx uses docx.Document(file) which accepts path or file-like
        
        # Mocking the file object structure if needed, but here we pass path or open file
        # streamlit file_uploader returns a file-like object with .type attribute
        
        class MockFile:
            def __init__(self, path, type):
                self.path = path
                self.type = type
                self._file = open(path, 'rb')
            
            def read(self):
                return self._file.read()
            
            def seek(self, offset):
                self._file.seek(offset)
            
            def close(self):
                self._file.close()

            # For docx.Document, it handles file path or stream. 
            # In ingestion.py: doc = docx.Document(file)
            # If we pass MockFile, docx might fail if it doesn't look like a stream.
            # But we can pass the path for direct testing of `extract_text_from_docx`
            # and a stream for `extract_resume_text`

        # Direct function test
        text = extract_text_from_docx(filename)
        self.assertIn("Hello World", text)
        self.assertIn("This is a resume", text)

        # Clean up
        os.remove(filename)

    @patch("core.database.get_document_collection")

    def test_save_document(self, mock_get_col):
        mock_col = MagicMock()
        mock_get_col.return_value = mock_col
        
        database.save_document("user@example.com", "Resume Text", "Job Description")

        
        mock_col.insert_one.assert_called_once()
        args = mock_col.insert_one.call_args[0][0]
        self.assertEqual(args["user_email"], "user@example.com")
        self.assertEqual(args["resume_text"], "Resume Text")
        self.assertEqual(args["job_description"], "Job Description")
        self.assertTrue("created_at" in args)

if __name__ == "__main__":
    unittest.main()
